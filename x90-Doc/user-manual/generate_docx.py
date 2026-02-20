#!/usr/bin/env python3
"""Generate official bilingual (EN/TH) Word document user manual for xMixing."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SS_DIR = os.path.join(SCRIPT_DIR, "screenshots")

# ─── Color Constants ───
BLUE_PRIMARY = RGBColor(21, 101, 192)    # #1565C0
BLUE_DARK    = RGBColor(13, 71, 161)     # #0D47A1
BLUE_LIGHT   = RGBColor(227, 242, 253)   # #E3F2FD
ORANGE       = RGBColor(230, 81, 0)      # #E65100
ORANGE_LIGHT = RGBColor(255, 243, 224)   # #FFF3E0
GRAY         = RGBColor(100, 100, 100)
WHITE        = RGBColor(255, 255, 255)
BLACK        = RGBColor(0, 0, 0)
GREEN_BG     = RGBColor(232, 245, 233)   # #E8F5E9

# ─── Helper Functions ───
def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_paragraph(doc, text, style=None, font_size=None, bold=False,
                            color=None, alignment=None, space_before=None,
                            space_after=None, font_name=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if font_size:
        run.font.size = Pt(font_size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet_list(doc, items, bold_prefix=True):
    """Add bullet list items. Each item can be 'Bold Part — rest' or plain text."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        if '—' in item and bold_prefix:
            parts = item.split('—', 1)
            run_bold = p.add_run(parts[0].strip())
            run_bold.bold = True
            run_bold.font.size = Pt(10.5)
            run_rest = p.add_run(' — ' + parts[1].strip())
            run_rest.font.size = Pt(10.5)
        else:
            run = p.add_run(item)
            run.font.size = Pt(10.5)

def add_numbered_list(doc, items):
    """Add numbered list items."""
    for item in items:
        p = doc.add_paragraph(style='List Number')
        if '—' in item:
            parts = item.split('—', 1)
            run_bold = p.add_run(parts[0].strip())
            run_bold.bold = True
            run_bold.font.size = Pt(10.5)
            run_rest = p.add_run(' — ' + parts[1].strip())
            run_rest.font.size = Pt(10.5)
        else:
            run = p.add_run(item)
            run.font.size = Pt(10.5)

def add_tip_box(doc, text, is_warning=False):
    """Add a colored tip/warning box using a single-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    icon = "⚠️" if is_warning else "💡"
    label = "Important" if is_warning else "Tip"
    run = p.add_run(f"{icon} {label}: ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    color_hex = "FFF3E0" if is_warning else "E8F5E9"
    set_cell_shading(cell, color_hex)
    # Set table width
    table.columns[0].width = Inches(6.5)
    doc.add_paragraph()  # spacer

def add_screenshot(doc, filename, caption=""):
    """Add a screenshot image with optional caption."""
    path = os.path.join(SS_DIR, filename)
    if not os.path.exists(path):
        add_formatted_paragraph(doc, f"[Image not found: {filename}]",
                                font_size=10, color=GRAY,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(6.2))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = cap.add_run(caption)
        run_cap.font.size = Pt(9)
        run_cap.font.color.rgb = GRAY
        run_cap.italic = True

def add_chapter_header(doc, num, en_title, th_title):
    """Add a styled chapter header using a table."""
    doc.add_page_break()
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    # English title
    p1 = cell.paragraphs[0]
    run1 = p1.add_run(f"Chapter {num}: {en_title}")
    run1.bold = True
    run1.font.size = Pt(20)
    run1.font.color.rgb = WHITE
    # Thai subtitle
    p2 = cell.add_paragraph()
    run2 = p2.add_run(f"บทที่ {num}: {th_title}")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(200, 220, 255)
    set_cell_shading(cell, "1565C0")
    table.columns[0].width = Inches(6.5)
    doc.add_paragraph()  # spacer

def add_bilingual_section(doc, en_title, en_content_fn, th_title, th_content_fn):
    """Add a two-column bilingual content section using a table."""
    # Create 2-column table
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # English column
    en_cell = table.cell(0, 0)
    en_cell.text = ""
    set_cell_shading(en_cell, "F5F9FF")
    p_label = en_cell.paragraphs[0]
    run_label = p_label.add_run("🇬🇧 ENGLISH")
    run_label.bold = True
    run_label.font.size = Pt(8)
    run_label.font.color.rgb = GRAY
    en_content_fn(en_cell)

    # Thai column
    th_cell = table.cell(0, 1)
    th_cell.text = ""
    set_cell_shading(th_cell, "FFF8F0")
    p_label2 = th_cell.paragraphs[0]
    run_label2 = p_label2.add_run("🇹🇭 ภาษาไทย")
    run_label2.bold = True
    run_label2.font.size = Pt(8)
    run_label2.font.color.rgb = GRAY
    th_content_fn(th_cell)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(3.25)
        row.cells[1].width = Inches(3.25)

    doc.add_paragraph()  # spacer

def add_cell_paragraph(cell, text, bold=False, font_size=10.5):
    """Add a paragraph inside a table cell."""
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    return p

def add_cell_bullets(cell, items):
    """Add bullet points inside a table cell."""
    for item in items:
        p = cell.add_paragraph()
        p.style = doc.styles['List Bullet']
        if '—' in item:
            parts = item.split('—', 1)
            r1 = p.add_run(f"• {parts[0].strip()}")
            r1.bold = True
            r1.font.size = Pt(9.5)
            r2 = p.add_run(f" — {parts[1].strip()}")
            r2.font.size = Pt(9.5)
        else:
            r = p.add_run(f"• {item}")
            r.font.size = Pt(9.5)

def add_cell_numbered(cell, items):
    """Add numbered items inside a table cell."""
    for i, item in enumerate(items, 1):
        p = cell.add_paragraph()
        if '—' in item:
            parts = item.split('—', 1)
            r1 = p.add_run(f"{i}. {parts[0].strip()}")
            r1.bold = True
            r1.font.size = Pt(9.5)
            r2 = p.add_run(f" — {parts[1].strip()}")
            r2.font.size = Pt(9.5)
        else:
            r = p.add_run(f"{i}. {item}")
            r.font.size = Pt(9.5)


# ═══════════════════════════════════════
#  MAIN DOCUMENT GENERATION
# ═══════════════════════════════════════
doc = Document()

# ─── Page Setup ───
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ─── Styles ───
style = doc.styles['Normal']
style.font.size = Pt(11)
style.font.name = 'Calibri'

# ═══════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

add_formatted_paragraph(doc, "🏭 xMixing", font_size=36, bold=True,
                        color=BLUE_PRIMARY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_after=6)
add_formatted_paragraph(doc, "User Manual / คู่มือผู้ใช้", font_size=22,
                        color=BLUE_DARK, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_after=20)

# Separator line
p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_line = p_line.add_run("━" * 50)
run_line.font.color.rgb = BLUE_PRIMARY
run_line.font.size = Pt(12)

add_formatted_paragraph(doc, "Batch Management & Mixing Control System",
                        font_size=14, color=GRAY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_formatted_paragraph(doc, "ระบบจัดการแบตช์และควบคุมการผสม",
                        font_size=13, color=GRAY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

doc.add_paragraph()
doc.add_paragraph()

add_formatted_paragraph(doc, "Version 1.0.0", font_size=12,
                        color=GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_formatted_paragraph(doc, "February 2026", font_size=12,
                        color=GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_formatted_paragraph(doc, "devTeam@xDev.co.th", font_size=11,
                        color=BLUE_PRIMARY, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
doc.add_paragraph()

# Confidential notice
add_formatted_paragraph(doc, "CONFIDENTIAL — FOR INTERNAL USE ONLY",
                        font_size=10, bold=True, color=GRAY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_formatted_paragraph(doc, "เอกสารลับ — สำหรับใช้ภายในเท่านั้น",
                        font_size=10, color=GRAY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)


# ═══════════════════════════════════════
#  TABLE OF CONTENTS
# ═══════════════════════════════════════
doc.add_page_break()
add_formatted_paragraph(doc, "📋 Table of Contents / สารบัญ", font_size=20,
                        bold=True, color=BLUE_PRIMARY, space_after=12)

# Separator
p_sep = doc.add_paragraph()
run_sep = p_sep.add_run("━" * 60)
run_sep.font.color.rgb = BLUE_PRIMARY
run_sep.font.size = Pt(8)

toc_items = [
    ("1", "Login", "เข้าสู่ระบบ"),
    ("2", "Dashboard", "แดชบอร์ด"),
    ("3", "Ingredient Intake", "รับวัตถุดิบ"),
    ("4", "Ingredient Configuration", "ตั้งค่าวัตถุดิบ"),
    ("5", "SKU Management", "จัดการ SKU"),
    ("6", "Production Plan", "แผนการผลิต"),
    ("7", "Batch Prepare (Pre-Batch Weighing)", "เตรียมแบตช์ (ชั่งน้ำหนัก Pre-Batch)"),
    ("8", "Packing List", "รายการบรรจุ"),
    ("9", "User Management (Admin)", "จัดการผู้ใช้ (แอดมิน)"),
    ("10", "System Dashboard (Admin)", "แดชบอร์ดระบบ (แอดมิน)"),
]

# TOC table
toc_table = doc.add_table(rows=len(toc_items), cols=3)
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (num, en, th) in enumerate(toc_items):
    cells = toc_table.rows[i].cells
    # Chapter number
    r = cells[0].paragraphs[0].add_run(f"Chapter {num}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BLUE_PRIMARY
    # English title
    r2 = cells[1].paragraphs[0].add_run(en)
    r2.font.size = Pt(11)
    # Thai title
    r3 = cells[2].paragraphs[0].add_run(th)
    r3.font.size = Pt(11)
    r3.font.color.rgb = GRAY
    # Alternate row shading
    if i % 2 == 0:
        for cell in cells:
            set_cell_shading(cell, "F5F9FF")

toc_table.columns[0].width = Inches(1.2)
toc_table.columns[1].width = Inches(2.8)
toc_table.columns[2].width = Inches(2.5)


# ═══════════════════════════════════════
#  CHAPTERS
# ═══════════════════════════════════════

# ─── Chapter 1: Login ───
add_chapter_header(doc, "1", "Login", "เข้าสู่ระบบ")
add_screenshot(doc, "00-login.png", "Login Page / หน้าเข้าสู่ระบบ")

add_formatted_paragraph(doc, "🇬🇧 English", font_size=9, bold=True, color=BLUE_PRIMARY, space_before=12)
add_formatted_paragraph(doc, "Open the application in a web browser. You will see the Login Page.", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "Steps:", font_size=11, bold=True, space_after=2)
add_numbered_list(doc, [
    "Enter your Username in the username field.",
    "Enter your Password in the password field.",
    "Click the \"Login\" button.",
    "If successful, you will be redirected to the Dashboard.",
])
add_tip_box(doc, "If you don't have an account, click \"Register here\" to create one.")

add_formatted_paragraph(doc, "🇹🇭 ภาษาไทย", font_size=9, bold=True, color=ORANGE, space_before=12)
add_formatted_paragraph(doc, "เปิดแอปพลิเคชันในเว็บเบราว์เซอร์ คุณจะเห็นหน้าเข้าสู่ระบบ", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "ขั้นตอน:", font_size=11, bold=True, space_after=2)
add_numbered_list(doc, [
    "กรอกชื่อผู้ใช้ในช่องชื่อผู้ใช้",
    "กรอกรหัสผ่านในช่องรหัสผ่าน",
    "คลิกปุ่ม \"เข้าสู่ระบบ\"",
    "หากสำเร็จ ระบบจะนำคุณไปยังแดชบอร์ด",
])
add_tip_box(doc, "หากยังไม่มีบัญชี ให้คลิก \"ลงทะเบียนที่นี่\" เพื่อสร้างบัญชีใหม่")


# ─── Chapter 2: Dashboard ───
add_chapter_header(doc, "2", "Dashboard", "แดชบอร์ด")
add_screenshot(doc, "01-dashboard.png", "Dashboard / แดชบอร์ด")

add_formatted_paragraph(doc, "🇬🇧 English", font_size=9, bold=True, color=BLUE_PRIMARY, space_before=12)
add_formatted_paragraph(doc, "The Dashboard is the main landing page after login. It provides an overview of the system status.", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "Key Elements:", font_size=11, bold=True, space_after=2)
add_bullet_list(doc, [
    "Welcome Banner — Shows your username, role, and account status.",
    "Statistics Cards — Displays Total SKUs, Ingredients Stock, Pending Batches, and Active Productions.",
    "Quick Access Buttons — Shortcuts to Create SKU, Ingredient Intake, Plan Batch, and Start Production.",
    "Recent Activities — Timeline of recent system events and batch updates.",
])

add_formatted_paragraph(doc, "🇹🇭 ภาษาไทย", font_size=9, bold=True, color=ORANGE, space_before=12)
add_formatted_paragraph(doc, "แดชบอร์ดเป็นหน้าหลักหลังจากเข้าสู่ระบบ แสดงภาพรวมสถานะของระบบ", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "องค์ประกอบหลัก:", font_size=11, bold=True, space_after=2)
add_bullet_list(doc, [
    "แบนเนอร์ต้อนรับ — แสดงชื่อผู้ใช้ บทบาท และสถานะบัญชี",
    "การ์ดสถิติ — แสดง SKU ทั้งหมด, สต็อกวัตถุดิบ, แบตช์ที่รอดำเนินการ, และการผลิตที่กำลังดำเนินอยู่",
    "ปุ่มทางลัด — ลัดไปยังสร้าง SKU, รับวัตถุดิบ, วางแผนแบตช์, และเริ่มการผลิต",
    "กิจกรรมล่าสุด — ไทม์ไลน์เหตุการณ์ล่าสุดในระบบ",
])


# ─── Chapter 3: Ingredient Intake ───
add_chapter_header(doc, "3", "Ingredient Intake", "รับวัตถุดิบ")
add_screenshot(doc, "02-ingredient-intake.png", "Ingredient Intake / รับวัตถุดิบ")

add_formatted_paragraph(doc, "🇬🇧 English", font_size=9, bold=True, color=BLUE_PRIMARY, space_before=12)
add_formatted_paragraph(doc, "The Ingredient Intake page is used to log incoming raw materials into the system.", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "Steps to Record an Intake:", font_size=11, bold=True, space_after=2)
add_numbered_list(doc, [
    "Scan or type the Ingredient Code — The system auto-fills MAT.SAP Code, Re-Code, and Ingredient Name.",
    "Select the Intake Warehouse Location from the dropdown.",
    "Enter the Lot Number and optionally a PO Number.",
    "Set Manufacturing Date and Expire Date.",
    "Enter Intake Volume (kg) and Package Volume (kg). Num of Packages is auto-calculated.",
    "Click \"Save Intake\" to record the entry.",
])

add_formatted_paragraph(doc, "🇹🇭 ภาษาไทย", font_size=9, bold=True, color=ORANGE, space_before=12)
add_formatted_paragraph(doc, "หน้ารับวัตถุดิบใช้สำหรับบันทึกวัตถุดิบที่เข้ามาในระบบ", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "ขั้นตอนการบันทึกการรับ:", font_size=11, bold=True, space_after=2)
add_numbered_list(doc, [
    "สแกนหรือพิมพ์รหัสวัตถุดิบ — ระบบจะกรอก MAT.SAP Code, Re-Code, และชื่อวัตถุดิบอัตโนมัติ",
    "เลือกสถานที่รับ (คลังสินค้า) จากเมนูดรอปดาวน์",
    "กรอกหมายเลข Lot และหมายเลข PO (ถ้ามี)",
    "ตั้งวันที่ผลิตและวันหมดอายุ",
    "กรอกปริมาณรับ (กก.) และปริมาณต่อถุง (กก.) จำนวนถุงจะคำนวณอัตโนมัติ",
    "คลิก \"บันทึกการรับ\" เพื่อบันทึกรายการ",
])


# ─── Chapter 4: Ingredient Configuration ───
add_chapter_header(doc, "4", "Ingredient Configuration", "ตั้งค่าวัตถุดิบ")
add_screenshot(doc, "03-ingredient-config.png", "Ingredient Configuration / ตั้งค่าวัตถุดิบ")

add_formatted_paragraph(doc, "🇬🇧 English", font_size=9, bold=True, color=BLUE_PRIMARY, space_before=12)
add_formatted_paragraph(doc, "Manage the master list of ingredients used across the system.", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "Features:", font_size=11, bold=True, space_after=2)
add_bullet_list(doc, [
    "View — All ingredients in a searchable, sortable table.",
    "Add New Ingredient — Click \"+\" button, fill in Ingredient ID, Name, MAT.SAP Code, Re-Code, Description, and UOM.",
    "Edit — Click the edit icon on any row to modify ingredient details.",
    "Delete — Click the delete icon to remove an ingredient (requires confirmation).",
    "Print Labels — Generate and print barcode labels for ingredients.",
])

add_formatted_paragraph(doc, "🇹🇭 ภาษาไทย", font_size=9, bold=True, color=ORANGE, space_before=12)
add_formatted_paragraph(doc, "จัดการรายการหลักของวัตถุดิบที่ใช้ทั่วทั้งระบบ", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "คุณสมบัติ:", font_size=11, bold=True, space_after=2)
add_bullet_list(doc, [
    "ดู — วัตถุดิบทั้งหมดในตารางที่ค้นหาและเรียงลำดับได้",
    "เพิ่มวัตถุดิบใหม่ — คลิกปุ่ม \"+\" กรอก ID วัตถุดิบ, ชื่อ, MAT.SAP Code, Re-Code, คำอธิบาย, และหน่วยวัด",
    "แก้ไข — คลิกไอคอนแก้ไขในแถวเพื่อแก้ไขรายละเอียด",
    "ลบ — คลิกไอคอนลบเพื่อลบวัตถุดิบ (ต้องยืนยัน)",
    "พิมพ์ฉลาก — สร้างและพิมพ์ฉลากบาร์โค้ดสำหรับวัตถุดิบ",
])


# ─── Chapter 5: SKU Management ───
add_chapter_header(doc, "5", "SKU Management", "จัดการ SKU")
add_screenshot(doc, "04-sku.png", "SKU Management / จัดการ SKU")

add_formatted_paragraph(doc, "🇬🇧 English", font_size=9, bold=True, color=BLUE_PRIMARY, space_before=12)
add_formatted_paragraph(doc, "The SKU Management page allows you to create and manage product recipes (SKUs).", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "Features:", font_size=11, bold=True, space_after=2)
add_bullet_list(doc, [
    "SKU List — View all existing SKUs with their details.",
    "Create New SKU — Define a new product with SKU ID, Name, Batch Size (kg), and Plant assignment.",
    "Recipe Builder — Add ingredients to an SKU with specific percentages or weights.",
    "Package Configuration — Define package types and sizes for each SKU.",
    "Edit / Delete — Modify or remove existing SKUs.",
])
add_tip_box(doc, "Ensure ingredient percentages total 100% for accurate batch calculations.")

add_formatted_paragraph(doc, "🇹🇭 ภาษาไทย", font_size=9, bold=True, color=ORANGE, space_before=12)
add_formatted_paragraph(doc, "หน้าจัดการ SKU ให้คุณสร้างและจัดการสูตรผลิตภัณฑ์ (SKU)", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "คุณสมบัติ:", font_size=11, bold=True, space_after=2)
add_bullet_list(doc, [
    "รายการ SKU — ดู SKU ทั้งหมดพร้อมรายละเอียด",
    "สร้าง SKU ใหม่ — กำหนดผลิตภัณฑ์ใหม่ด้วย SKU ID, ชื่อ, ขนาดแบตช์ (กก.), และโรงงาน",
    "ตัวสร้างสูตร — เพิ่มวัตถุดิบใน SKU พร้อมเปอร์เซ็นต์หรือน้ำหนักเฉพาะ",
    "การตั้งค่าบรรจุภัณฑ์ — กำหนดประเภทและขนาดบรรจุภัณฑ์สำหรับแต่ละ SKU",
    "แก้ไข / ลบ — แก้ไขหรือลบ SKU ที่มีอยู่",
])
add_tip_box(doc, "ตรวจสอบให้แน่ใจว่าเปอร์เซ็นต์วัตถุดิบรวมเป็น 100% เพื่อการคำนวณแบตช์ที่แม่นยำ")


# ─── Chapter 6: Production Plan ───
add_chapter_header(doc, "6", "Production Plan", "แผนการผลิต")
add_screenshot(doc, "05-production-plan.png", "Production Plan / แผนการผลิต")

add_formatted_paragraph(doc, "🇬🇧 English", font_size=9, bold=True, color=BLUE_PRIMARY, space_before=12)
add_formatted_paragraph(doc, "The Production Plan page is used to schedule and organize production runs.", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "Steps to Create a Plan:", font_size=11, bold=True, space_after=2)
add_numbered_list(doc, [
    "Select a SKU from the dropdown.",
    "Enter the Total Target Volume (kg).",
    "Select the Plant/Production Line.",
    "Click \"Create Plan\" — The system auto-calculates the number of batches required based on the plant's batch capacity.",
])
add_tip_box(doc, "Plans can be printed for distribution to the production floor.")

add_formatted_paragraph(doc, "🇹🇭 ภาษาไทย", font_size=9, bold=True, color=ORANGE, space_before=12)
add_formatted_paragraph(doc, "หน้าแผนการผลิตใช้สำหรับจัดตารางและจัดระเบียบรอบการผลิต", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "ขั้นตอนการสร้างแผน:", font_size=11, bold=True, space_after=2)
add_numbered_list(doc, [
    "เลือก SKU จากเมนูดรอปดาวน์",
    "กรอกปริมาณเป้าหมายรวม (กก.)",
    "เลือกโรงงาน/สายการผลิต",
    "คลิก \"สร้างแผน\" — ระบบจะคำนวณจำนวนแบตช์ที่ต้องการโดยอัตโนมัติ",
])
add_tip_box(doc, "สามารถพิมพ์แผนเพื่อแจกจ่ายไปยังพื้นที่การผลิตได้")


# ─── Chapter 7: Batch Prepare ───
add_chapter_header(doc, "7", "Batch Prepare (Pre-Batch Weighing)", "เตรียมแบตช์ (ชั่งน้ำหนัก Pre-Batch)")
add_screenshot(doc, "06-pre-batch.png", "Batch Prepare / เตรียมแบตช์")

add_formatted_paragraph(doc, "🇬🇧 English", font_size=9, bold=True, color=BLUE_PRIMARY, space_before=12)
add_formatted_paragraph(doc, "The Batch Prepare page is the core operational screen where operators weigh ingredients for each batch.", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "Workflow:", font_size=11, bold=True, space_after=2)
add_numbered_list(doc, [
    "Select a Production Plan from the list.",
    "Select a specific Batch to work on.",
    "The system displays each ingredient required with its target volume.",
    "Scan the ingredient barcode to identify the material.",
    "Place the ingredient on the scale — The system reads the weight in real-time via MQTT integration.",
    "Confirm the weight and move to the next ingredient.",
])

add_formatted_paragraph(doc, "🇹🇭 ภาษาไทย", font_size=9, bold=True, color=ORANGE, space_before=12)
add_formatted_paragraph(doc, "หน้าเตรียมแบตช์เป็นหน้าจอปฏิบัติการหลักที่พนักงานชั่งวัตถุดิบสำหรับแต่ละแบตช์", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "ขั้นตอนการทำงาน:", font_size=11, bold=True, space_after=2)
add_numbered_list(doc, [
    "เลือกแผนการผลิตจากรายการ",
    "เลือกแบตช์เฉพาะที่จะทำงาน",
    "ระบบจะแสดงวัตถุดิบแต่ละรายการที่ต้องการพร้อมปริมาณเป้าหมาย",
    "สแกนบาร์โค้ดวัตถุดิบเพื่อระบุวัสดุ",
    "วางวัตถุดิบบนเครื่องชั่ง — ระบบอ่านน้ำหนักแบบเรียลไทม์ผ่าน MQTT",
    "ยืนยันน้ำหนักและไปยังวัตถุดิบถัดไป",
])


# ─── Chapter 8: Packing List ───
add_chapter_header(doc, "8", "Packing List", "รายการบรรจุ")
add_screenshot(doc, "07-packing-list.png", "Packing List / รายการบรรจุ")

add_formatted_paragraph(doc, "🇬🇧 English", font_size=9, bold=True, color=BLUE_PRIMARY, space_before=12)
add_formatted_paragraph(doc, "The Packing List page manages the final boxing and verification process.", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "Features:", font_size=11, bold=True, space_after=2)
add_bullet_list(doc, [
    "Production Plan List — View plans with batch IDs, SKUs, volumes, and pack counts.",
    "2-Step Verification — Scan an ingredient bag barcode, then scan the Box ID to confirm.",
    "Confirm Packing Table — Finalize and save the packing configuration.",
    "Print List — Queue box labels for printing with all ingredient details.",
    "Pre-Batch Scans — View all scans associated with a selected batch.",
])
add_tip_box(doc, "Always verify scans before confirming the packing table.", is_warning=True)

add_formatted_paragraph(doc, "🇹🇭 ภาษาไทย", font_size=9, bold=True, color=ORANGE, space_before=12)
add_formatted_paragraph(doc, "หน้ารายการบรรจุจัดการกระบวนการบรรจุกล่องและตรวจสอบขั้นสุดท้าย", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "คุณสมบัติ:", font_size=11, bold=True, space_after=2)
add_bullet_list(doc, [
    "รายการแผนการผลิต — ดูแผนพร้อม Batch ID, SKU, ปริมาณ, และจำนวนแพ็ค",
    "การตรวจสอบ 2 ขั้นตอน — สแกนบาร์โค้ดถุง จากนั้นสแกน Box ID เพื่อยืนยัน",
    "ยืนยันตาราง Packing — สรุปและบันทึกการตั้งค่าการบรรจุ",
    "รายการพิมพ์ — คิวฉลากกล่องสำหรับพิมพ์พร้อมรายละเอียดวัตถุดิบทั้งหมด",
    "รายการสแกน Pre-Batch — ดูการสแกนทั้งหมดที่เกี่ยวข้องกับแบตช์ที่เลือก",
])
add_tip_box(doc, "ตรวจสอบการสแกนก่อนยืนยันตาราง Packing เสมอ", is_warning=True)


# ─── Chapter 9: User Management ───
add_chapter_header(doc, "9", "User Management (Admin)", "จัดการผู้ใช้ (แอดมิน)")
add_screenshot(doc, "09-user-config.png", "User Management / จัดการผู้ใช้")

add_formatted_paragraph(doc, "🇬🇧 English", font_size=9, bold=True, color=BLUE_PRIMARY, space_before=12)
add_formatted_paragraph(doc, "The User Management page is for administrators to manage user accounts and permissions.", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "Features:", font_size=11, bold=True, space_after=2)
add_bullet_list(doc, [
    "User List — Search and view all users with Name, Email, Role, Department.",
    "Add User — Create new accounts with username, email, password, role, and department.",
    "Manage User — Click \"Manage\" to edit user info, change password, and configure permissions.",
    "Permissions — Toggle individual page access for each user.",
    "Delete User — Remove user accounts (requires confirmation).",
])

add_formatted_paragraph(doc, "🇹🇭 ภาษาไทย", font_size=9, bold=True, color=ORANGE, space_before=12)
add_formatted_paragraph(doc, "หน้าจัดการผู้ใช้สำหรับผู้ดูแลระบบเพื่อจัดการบัญชีผู้ใช้และสิทธิ์", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "คุณสมบัติ:", font_size=11, bold=True, space_after=2)
add_bullet_list(doc, [
    "รายการผู้ใช้ — ค้นหาและดูผู้ใช้ทั้งหมดพร้อมชื่อ อีเมล บทบาท แผนก",
    "เพิ่มผู้ใช้ — สร้างบัญชีผู้ใช้ใหม่ด้วยชื่อผู้ใช้ อีเมล รหัสผ่าน บทบาท และแผนก",
    "จัดการผู้ใช้ — คลิก \"จัดการ\" เพื่อแก้ไขข้อมูล เปลี่ยนรหัสผ่าน และตั้งค่าสิทธิ์",
    "สิทธิ์ — สลับการเข้าถึงหน้าแต่ละหน้าสำหรับผู้ใช้แต่ละคน",
    "ลบผู้ใช้ — ลบบัญชีผู้ใช้ (ต้องยืนยัน)",
])


# ─── Chapter 10: System Dashboard ───
add_chapter_header(doc, "10", "System Dashboard (Admin)", "แดชบอร์ดระบบ (แอดมิน)")
add_screenshot(doc, "10-system-dashboard.png", "System Dashboard / แดชบอร์ดระบบ")

add_formatted_paragraph(doc, "🇬🇧 English", font_size=9, bold=True, color=BLUE_PRIMARY, space_before=12)
add_formatted_paragraph(doc, "The System Dashboard provides real-time monitoring of the server infrastructure.", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "Metrics Displayed:", font_size=11, bold=True, space_after=2)
add_bullet_list(doc, [
    "PC Information — Hostname, IP Address, OS, Architecture, CPU Model.",
    "System Uptime — Boot time and uptime duration.",
    "CPU Usage — Real-time CPU utilization with circular gauge.",
    "Memory (RAM) — Memory usage with used/total display.",
    "Storage (Disk) — Disk usage with used/total display.",
    "Network Traffic — Bytes sent and received.",
    "History Charts — CPU and Memory usage over the last 1 hour.",
])

add_formatted_paragraph(doc, "🇹🇭 ภาษาไทย", font_size=9, bold=True, color=ORANGE, space_before=12)
add_formatted_paragraph(doc, "แดชบอร์ดระบบให้การตรวจสอบโครงสร้างพื้นฐานของเซิร์ฟเวอร์แบบเรียลไทม์", font_size=10.5, space_after=4)
add_formatted_paragraph(doc, "เมตริกที่แสดง:", font_size=11, bold=True, space_after=2)
add_bullet_list(doc, [
    "ข้อมูล PC — ชื่อโฮสต์, ที่อยู่ IP, OS, สถาปัตยกรรม, รุ่น CPU",
    "เวลาทำงานของระบบ — เวลาบูตและระยะเวลาทำงาน",
    "การใช้ CPU — เปอร์เซ็นต์การใช้ CPU แบบเรียลไทม์พร้อมมาตรวัดวงกลม",
    "หน่วยความจำ (RAM) — การใช้หน่วยความจำพร้อมแสดงใช้แล้ว/ทั้งหมด",
    "พื้นที่จัดเก็บ (ดิสก์) — การใช้ดิสก์พร้อมแสดงใช้แล้ว/ทั้งหมด",
    "ทราฟฟิกเครือข่าย — ไบต์ที่ส่งและรับ",
    "กราฟประวัติ — การใช้ CPU และหน่วยความจำในช่วง 1 ชั่วโมงที่ผ่านมา",
])


# ═══════════════════════════════════════
#  FOOTER / BACK PAGE
# ═══════════════════════════════════════
doc.add_page_break()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
add_formatted_paragraph(doc, "━" * 50, font_size=10, color=BLUE_PRIMARY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_formatted_paragraph(doc, "© 2026 xMixing — All rights reserved.",
                        font_size=12, color=GRAY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_formatted_paragraph(doc, "Developed by devTeam@xDev.co.th",
                        font_size=11, color=GRAY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_formatted_paragraph(doc, "พัฒนาโดย devTeam@xDev.co.th",
                        font_size=11, color=GRAY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
add_formatted_paragraph(doc, "For support, contact: devTeam@xDev.co.th",
                        font_size=10, color=BLUE_PRIMARY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_formatted_paragraph(doc, "สำหรับการสนับสนุน ติดต่อ: devTeam@xDev.co.th",
                        font_size=10, color=ORANGE,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)


# ─── Save Document ───
output_path = os.path.join(SCRIPT_DIR, "xMixing-UserManual.docx")
doc.save(output_path)

file_size_mb = os.path.getsize(output_path) / 1024 / 1024
print(f"✅ Word document generated: {output_path}")
print(f"   File size: {file_size_mb:.1f} MB")
print(f"   Chapters: 10")
print(f"   Languages: English + Thai")
print(f"\n📄 Open with: open \"{output_path}\"")
